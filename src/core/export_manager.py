"""
导出管理器
负责将项目导出为各种格式（文本、Word、PDF、HTML等）
"""

import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass
from enum import Enum

from PyQt6.QtCore import QObject, pyqtSignal, QThread

from .concurrent_io import FileIOWorker, get_concurrent_io

from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from core.project import ProjectManager, ProjectDocument, DocumentType

logger = logging.getLogger(__name__)


class ExportFormat(Enum):
    """导出格式枚举"""
    TEXT = "text"
    MARKDOWN = "markdown"
    DOCX = "docx"
    PDF = "pdf"
    HTML = "html"
    EPUB = "epub"


@dataclass
class ExportOptions:
    """导出选项"""
    format: ExportFormat
    output_path: Path
    include_metadata: bool = True
    include_comments: bool = False
    preserve_formatting: bool = True
    chapter_break: str = "\n\n---\n\n"  # 章节分隔符
    encoding: str = "utf-8"
    title: Optional[str] = None
    author: Optional[str] = None


class ExportManager(QObject):
    """导出管理器"""
    
    # 信号
    exportStarted = pyqtSignal(str)  # 导出开始
    exportProgress = pyqtSignal(int, int)  # 当前进度, 总数
    exportCompleted = pyqtSignal(str)  # 导出完成
    exportError = pyqtSignal(str)  # 导出错误
    
    def __init__(self, project_manager: 'ProjectManager'):
        super().__init__()
        self._project_manager = project_manager
        
    def export_project(self, options: ExportOptions) -> bool:
        """导出项目"""
        try:
            self.exportStarted.emit(f"开始导出为 {options.format.value} 格式...")
            
            # 获取当前项目
            project = self._project_manager.get_current_project()
            if not project:
                self.exportError.emit("没有打开的项目")
                return False
            
            # 根据格式选择导出方法
            if options.format == ExportFormat.TEXT:
                return self._export_to_text(project, options)
            elif options.format == ExportFormat.MARKDOWN:
                return self._export_to_markdown(project, options)
            elif options.format == ExportFormat.DOCX:
                return self._export_to_docx(project, options)
            elif options.format == ExportFormat.PDF:
                return self._export_to_pdf(project, options)
            elif options.format == ExportFormat.HTML:
                return self._export_to_html(project, options)
            else:
                self.exportError.emit(f"不支持的导出格式: {options.format.value}")
                return False
                
        except Exception as e:
            logger.error(f"导出失败: {e}")
            self.exportError.emit(str(e))
            return False
    
    def _collect_documents(self, project: Any) -> List['ProjectDocument']:
        """收集要导出的文档（按顺序）"""
        documents = []
        doc_dict = project.documents
        
        # 构建父子关系映射
        children_map = {}
        root_docs = []
        
        for doc_id, doc in doc_dict.items():
            # 只导出小说内容
            if doc.doc_type.value in ['act', 'chapter', 'scene']:
                if doc.parent_id:
                    if doc.parent_id not in children_map:
                        children_map[doc.parent_id] = []
                    children_map[doc.parent_id].append(doc)
                else:
                    root_docs.append(doc)
        
        # 递归收集文档
        def collect_recursive(doc_list: List, parent_id: Optional[str] = None):
            for doc in sorted(doc_list, key=lambda d: d.order):
                documents.append(doc)
                if doc.id in children_map:
                    collect_recursive(children_map[doc.id], doc.id)
        
        collect_recursive(root_docs)
        return documents
    
    def _export_to_text(self, project: Any, options: ExportOptions) -> bool:
        """导出为纯文本（非阻塞版本）"""
        try:
            documents = self._collect_documents(project)
            total = len(documents)
            
            # 准备内容
            content_parts = []
            
            # 添加标题和作者信息
            if options.include_metadata:
                title = options.title or project.name
                author = options.author or project.author
                content_parts.append(f"{title}\n")
                content_parts.append(f"作者：{author}\n")
                content_parts.append("\n" + "="*50 + "\n\n")
            
            # 生成文档内容
            for i, doc in enumerate(documents):
                self.exportProgress.emit(i + 1, total)
                
                # 添加章节标题
                if doc.doc_type.value == 'act':
                    content_parts.append(f"\n第{doc.order + 1}幕 {doc.name}\n")
                    content_parts.append("="*30 + "\n\n")
                elif doc.doc_type.value == 'chapter':
                    content_parts.append(f"\n第{doc.order + 1}章 {doc.name}\n")
                    content_parts.append("-"*30 + "\n\n")
                elif doc.doc_type.value == 'scene':
                    content_parts.append(f"\n场景{doc.order + 1}：{doc.name}\n\n")
                
                # 添加内容
                if doc.content:
                    content_parts.append(doc.content)
                    content_parts.append("\n")
                
                # 章节分隔
                if doc.doc_type.value in ['act', 'chapter']:
                    content_parts.append(options.chapter_break)
            
            # 合并所有内容
            full_content = ''.join(content_parts)
            
            # 使用非阻塞写入
            worker = FileIOWorker('write', options.output_path, 
                                data=full_content, encoding=options.encoding, 
                                parent=self)
            worker.finished.connect(lambda: self.exportCompleted.emit(str(options.output_path)))
            worker.error.connect(lambda e: self.exportError.emit(f"导出文本失败: {e}"))
            worker.start()
            
            return True
            
        except Exception as e:
            logger.error(f"导出文本失败: {e}")
            self.exportError.emit(f"导出文本失败: {e}")
            return False
    
    def _export_to_markdown(self, project: Any, options: ExportOptions) -> bool:
        """导出为Markdown格式"""
        try:
            documents = self._collect_documents(project)
            total = len(documents)
            
            with open(options.output_path, 'w', encoding=options.encoding) as f:
                # 写入元数据
                if options.include_metadata:
                    title = options.title or project.name
                    author = options.author or project.author
                    f.write(f"# {title}\n\n")
                    f.write(f"**作者**: {author}\n\n")
                    f.write("---\n\n")
                
                # 写入文档内容
                for i, doc in enumerate(documents):
                    self.exportProgress.emit(i + 1, total)
                    
                    # 写入标题（使用Markdown标题级别）
                    if doc.doc_type.value == 'act':
                        f.write(f"\n# 第{doc.order + 1}幕 {doc.name}\n\n")
                    elif doc.doc_type.value == 'chapter':
                        f.write(f"\n## 第{doc.order + 1}章 {doc.name}\n\n")
                    elif doc.doc_type.value == 'scene':
                        f.write(f"\n### 场景{doc.order + 1}：{doc.name}\n\n")
                    
                    # 写入内容
                    if doc.content:
                        # 处理特殊标记
                        content = doc.content
                        # 保留@标记
                        content = content.replace('@', '**@') 
                        content = content.replace('**@', '@')
                        f.write(content)
                        f.write("\n\n")
            
            self.exportCompleted.emit(str(options.output_path))
            return True
            
        except Exception as e:
            logger.error(f"导出Markdown失败: {e}")
            self.exportError.emit(f"导出Markdown失败: {e}")
            return False
    
    def _export_to_docx(self, project: Any, options: ExportOptions) -> bool:
        """导出为Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            self.exportError.emit("需要安装python-docx库: pip install python-docx")
            return False
        
        try:
            documents = self._collect_documents(project)
            total = len(documents)
            
            # 创建Word文档
            doc = Document()
            
            # 设置标题和作者
            if options.include_metadata:
                title = options.title or project.name
                author = options.author or project.author
                
                # 标题
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_para.add_run(title)
                title_run.font.size = Pt(24)
                title_run.bold = True
                
                # 作者
                author_para = doc.add_paragraph()
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                author_run = author_para.add_run(f"作者：{author}")
                author_run.font.size = Pt(14)
                
                # 分页
                doc.add_page_break()
            
            # 写入文档内容
            for i, document in enumerate(documents):
                self.exportProgress.emit(i + 1, total)
                
                # 添加标题
                if document.doc_type.value == 'act':
                    heading = doc.add_heading(f"第{document.order + 1}幕 {document.name}", level=1)
                elif document.doc_type.value == 'chapter':
                    heading = doc.add_heading(f"第{document.order + 1}章 {document.name}", level=2)
                elif document.doc_type.value == 'scene':
                    heading = doc.add_heading(f"场景{document.order + 1}：{document.name}", level=3)
                
                # 添加内容
                if document.content:
                    # 按段落分割
                    paragraphs = document.content.split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text)
                            para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
                
                # 章节之间添加分页
                if document.doc_type.value in ['act', 'chapter'] and i < len(documents) - 1:
                    doc.add_page_break()
            
            # 保存文档
            doc.save(str(options.output_path))
            
            self.exportCompleted.emit(str(options.output_path))
            return True
            
        except Exception as e:
            logger.error(f"导出Word文档失败: {e}")
            self.exportError.emit(f"导出Word文档失败: {e}")
            return False
    
    def _export_to_pdf(self, project: Any, options: ExportOptions) -> bool:
        """导出为PDF（通过HTML转换）"""
        try:
            # 先导出为HTML
            html_path = options.output_path.with_suffix('.html')
            html_options = ExportOptions(
                format=ExportFormat.HTML,
                output_path=html_path,
                include_metadata=options.include_metadata,
                title=options.title,
                author=options.author
            )
            
            if not self._export_to_html(project, html_options):
                return False
            
            # 使用weasyprint转换为PDF
            try:
                from weasyprint import HTML
                HTML(filename=str(html_path)).write_pdf(str(options.output_path))
                
                # 删除临时HTML文件
                if html_path.exists():
                    html_path.unlink()
                
                self.exportCompleted.emit(str(options.output_path))
                return True
                
            except ImportError:
                # 清理临时文件
                if html_path.exists():
                    html_path.unlink()
                self.exportError.emit("需要安装weasyprint库: pip install weasyprint")
                return False
            except Exception as weasy_error:
                # 清理临时文件
                if html_path.exists():
                    html_path.unlink()
                self.exportError.emit(f"PDF转换失败: {weasy_error}")
                return False
                
        except Exception as e:
            logger.error(f"导出PDF失败: {e}")
            self.exportError.emit(f"导出PDF失败: {e}")
            return False
    
    def _export_to_html(self, project: Any, options: ExportOptions) -> bool:
        """导出为HTML"""
        try:
            documents = self._collect_documents(project)
            total = len(documents)
            
            with open(options.output_path, 'w', encoding=options.encoding) as f:
                # HTML头部
                f.write("""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
""")
                
                title = options.title or project.name
                f.write(f"    <title>{title}</title>\n")
                
                # 添加样式
                f.write("""    <style>
        body {
            font-family: "Microsoft YaHei", "SimSun", serif;
            line-height: 1.8;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }
        .content {
            background-color: white;
            padding: 40px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        h1 { text-align: center; margin-bottom: 30px; }
        h2 { margin-top: 40px; margin-bottom: 20px; }
        h3 { margin-top: 30px; margin-bottom: 15px; }
        p { text-indent: 2em; margin: 10px 0; }
        .author { text-align: center; font-size: 18px; margin-bottom: 50px; }
        .chapter-break { margin: 50px 0; text-align: center; }
    </style>
</head>
<body>
    <div class="content">
""")
                
                # 标题和作者
                if options.include_metadata:
                    author = options.author or project.author
                    f.write(f"        <h1>{title}</h1>\n")
                    f.write(f"        <p class='author'>作者：{author}</p>\n")
                
                # 写入内容
                for i, doc in enumerate(documents):
                    self.exportProgress.emit(i + 1, total)
                    
                    # 写入标题
                    if doc.doc_type.value == 'act':
                        f.write(f"        <h1>第{doc.order + 1}幕 {doc.name}</h1>\n")
                    elif doc.doc_type.value == 'chapter':
                        f.write(f"        <h2>第{doc.order + 1}章 {doc.name}</h2>\n")
                    elif doc.doc_type.value == 'scene':
                        f.write(f"        <h3>场景{doc.order + 1}：{doc.name}</h3>\n")
                    
                    # 写入内容
                    if doc.content:
                        paragraphs = doc.content.split('\n')
                        for para in paragraphs:
                            if para.strip():
                                # 转义HTML字符
                                para = para.replace('&', '&amp;')
                                para = para.replace('<', '&lt;')
                                para = para.replace('>', '&gt;')
                                f.write(f"        <p>{para}</p>\n")
                    
                    # 章节分隔
                    if doc.doc_type.value in ['act', 'chapter'] and i < len(documents) - 1:
                        f.write("        <div class='chapter-break'>* * *</div>\n")
                
                # HTML结尾
                f.write("""    </div>
</body>
</html>""")
            
            self.exportCompleted.emit(str(options.output_path))
            return True
            
        except Exception as e:
            logger.error(f"导出HTML失败: {e}")
            self.exportError.emit(f"导出HTML失败: {e}")
            return False